from __future__ import annotations

from io import BytesIO
from typing import Tuple

from docx import Document

from esg_tool.models.documents import ProcessDocument
from esg_tool.models.reports import ESGReportPackage


MIME_TYPE_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _sanitize_filename(name: str) -> str:
    """Return a filesystem friendly version of ``name`` without an extension."""
    cleaned = [c if c.isalnum() else "-" for c in name.strip() or "document"]
    filename = "".join(cleaned).strip("-") or "document"
    return filename


def _document_from_text(title: str, body: str) -> bytes:
    document = Document()
    if title:
        document.add_heading(title, level=1)
    for paragraph in (body or "").splitlines():
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def compiled_report_to_docx(report_package: ESGReportPackage) -> Tuple[str, bytes]:
    """Create a Word document for the compiled ESG report."""
    title = report_package.name or "ESG Report"
    filename = f"{_sanitize_filename(title)}.docx"
    binary = _document_from_text(title, report_package.compiled_report)
    return filename, binary


def process_document_to_docx(document: ProcessDocument) -> Tuple[str, bytes]:
    """Create a Word document from an intermediate process document."""
    title = document.title or "Process Document"
    filename = f"{_sanitize_filename(title)}.docx"
    binary = _document_from_text(title, document.content)
    return filename, binary
